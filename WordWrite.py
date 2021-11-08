from docx import Document


def write_docx(tasks_dict, docx_path):
    tasks_document = Document()

    for key_variant in tasks_dict.keys():
        if key_variant == 0:
            tasks_document.add_heading('Вариант для пересдачи')
        else:
            tasks_document.add_heading(f'Вариант {key_variant}')

        for n, task_dict in enumerate(tasks_dict[key_variant]):
            tasks_document.add_paragraph(
                f'{n + 1}. {task_dict["task"]}'
            )

    tasks_document.add_page_break()

    tasks_document.add_heading('Ответы', level=1)

    for key_variant in tasks_dict.keys():
        if key_variant == 0:
            tasks_document.add_heading('Вариант для пересдачи')
        else:
            tasks_document.add_heading(f'Вариант {key_variant}')

        for n, task_dict in enumerate(tasks_dict[key_variant]):
            tasks_document.add_paragraph(
                f'{n + 1}. {task_dict["answer"]}'
            )

    tasks_document.save(docx_path)


if __name__ == '__main__':
    demo_data = {
        0: [
            {'answer': '2', 'number': 1, 'task': '1 + 1'},
            {'answer': '2', 'number': 2, 'task': '1 * 2'},
            {'answer': '16', 'number': 3, 'task': '4 * 4'}
        ],
        1: [
            {'answer': '2', 'number': 1, 'task': '1 + 1'},
            {'answer': '2', 'number': 2, 'task': '1 * 2'},
            {'answer': '16', 'number': 3, 'task': '4 * 4'}],
        2: [
            {'answer': '2', 'number': 1, 'task': '1 + 1'},
            {'answer': '2', 'number': 2, 'task': '1 * 2'},
            {'answer': '16', 'number': 3, 'task': '4 * 4'}
        ],
        3: [
            {'answer': '2', 'number': 1, 'task': '1 + 1'},
            {'answer': '2', 'number': 2, 'task': '1 * 2'},
            {'answer': '16', 'number': 3, 'task': '4 * 4'}
        ]
    }
    write_docx(demo_data, 'test.docx')
